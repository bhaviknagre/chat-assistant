import os 
from typing import List
import pdfplumber 
from langchain.document_loaders import TextLoader
from langchain.docstore.document import Document
from docx import Document as DocxDocument
from utils.logger import logger 

def load_pdf(file_path: str) -> List[Document]:
    try:
        documents = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    documents.append(Document(
                        page_content = text,
                        metadata={
                            "source": file_path,
                            "page": page_num+1
                        }
                    ))
        logger.info(f"[PDF] Extracted {len(documents)} pages from {file_path}")
        return documents
    except Exception as e:
        logger.error(f"[PDF] Failed to extract pages from {file_path}: {e}")
        raise RuntimeError(f"PDF loading failed for {file_path}") from e
    
def load_txt(file_path: str) -> List[Document]:
    try:
        loader = TextLoader(file_path, encoding='utf-8')
        docs = loader.load()
        logger.info(f"[TXT Loader] Loaded text file: {file_path}")
        return docs
    except Exception as e:
        logger.error(f"[TXT Loader] Failed to load text file {file_path}: {e}")
        raise RuntimeError(f"TXT loading failed for {file_path}")
    
def load_docs(file_path: str) -> List[Document]:
    try:
        docx_obj = DocxDocument(file_path)
        full_text = "\n".join([para.text.strip() for para in docx_obj.paragraphs if para.text.strip()])
        metadata = {
            "source":file_path
        }
        document = Document(page_content=full_text, metadata=metadata)
        logger.info(f"[DOCX Loader] Loaded DOCX with {len(full_text.split())} words from {file_path}")
        return [document]
    except Exception as e:
        logger.error(f"[DOCX Loader] Error loading DOCX: {e}")
        raise RuntimeError(f"DOCX loading failed for {file_path}") from e
    
    
def load_document(file_path: str) -> List[Document]:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}") 
    
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return load_pdf(file_path)
    elif ext == '.txt':
        return load_txt(file_path)
    elif ext == '.docx':
        return load_docs(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .txt, .docx")